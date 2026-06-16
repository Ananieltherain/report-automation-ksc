"""
lhs_engine.py — LHS Generator Engine (v6)
PT. Karya Solusi Cemerlang

FIX v6:
  1. SPACING ANALISA HILANG
     Root cause v5: _remove_empty_analisa() menghapus semua paragraf kosong
     di zona analisa — padahal paragraf kosong antara [analisaN] dan [analisaN+1]
     adalah SPACING INTENTIONAL di template, bukan sisa yang perlu dihapus.
     Pola template: [analisaN] → '' (spacer) → [analisaN+1]
     Fix: hapus PASANGAN ([analisaN] + spacer-nya) sekaligus jika analisa tidak terpakai.
     Caranya: identifikasi index paragraf [analisaN] yang kosong, hapus dia DAN
     paragraf kosong tepat sesudahnya (spacer-nya).

  2. NEWLINE DARI INPUT TIDAK MUNCUL DI WORD
     Root cause: keterangan saksi & kronologis dikirim dari browser sebagai string
     dengan '\n' (newline Unix). Word tidak mengerti '\n' di dalam satu w:t —
     di Word, baris baru dalam satu paragraf = <w:br/> (line break) atau
     paragraf baru = <w:p> baru.
     Fix: saat replace keterangan/kronologis, split by '\n', buat multiple <w:r><w:br/>
     di dalam satu w:p (soft return, mempertahankan formatting paragraf).

  3. SPACING SAKSI HILANG — sama penyebabnya dengan analisa:
     _remove_orphan_empty_paragraphs() terlalu agresif menghapus empty di zona saksi.
     Fix: hapus paragraf KHUSUS yang menjadi penanda blok saksi tidak terpakai
     (nama+role kosong "( )", "Alamat :", usia kosong "(th)"), bukan hapus
     semua empty. Spacer empty antara blok saksi yang terpakai JANGAN disentuh.

  4. TTD SURVEYOR — include tanda tangan per surveyor
     Struktur template: "Dharmawan Gunawan, S.H. [surveyor]"
     [surveyor] diganti nama surveyor → menjadi:
     "Dharmawan Gunawan, S.H. Ahmad Fauzi"
     Untuk fitur TTD gambar: simpan file TTD per surveyor di folder 'ttd/'
     dengan nama file = nama surveyor + .png (contoh: "Ahmad Fauzi.png")
     Engine akan insert gambar TTD di atas baris nama jika file ditemukan.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, re, os

NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"

import pytesseract

# Sesuaikan jalur ini dengan lokasi instalasi Tesseract Anda
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Sekarang Anda bisa menggunakannya untuk OCR
# teks = pytesseract.image_to_string('gambar.png')

# ──────────────────────────────────────────────────────────────
# DATA REFERENSI
# ──────────────────────────────────────────────────────────────

DAFTAR_ASURANSI = [
    "PT. SOMPO INSURANCE INDONESIA",
    "PT. ZURICH INSURANCE",
    "PT. TOTAL ASURANSI BERSAMA",
    "PT. ASURANSI CENTRAL ASIA",
    "PT. MANDALA ARTHA GUNA",
]

DAFTAR_SURVEYOR = [
    "Albertus S.",
    "Yulius Bintoro",
    "Heru Prasetiyo",
    "Mawar Dinur Pane",
    "Arjianto Eko",
    "Samino",
    "Murtopo",
    "Freedy Setiawan",
]

DAFTAR_USAGE = [
    "Pribadi", "Komersial", "Dinas / Pemerintahan",
    "Rental / Sewa", "Taksi Online", "Angkutan Umum",
    "Operasional Perusahaan", "",
]

DAFTAR_COVERAGE = [
    "TLO (Total Loss Only)", "Comprehensive", "All Risk",
    "All Risk + Perluasan Banjir", "All Risk + Perluasan Gempa Bumi",
    "All Risk + Perluasan RSCC", "All Risk + Perluasan Lengkap", "",
]

DAFTAR_CASE = [
    "KEHILANGAN", "KERUSAKAN", "KERUSUHAN", "KEBAKARAN", "BANJIR", "",
]

PASAL_PSAKBI = [
    ("Pasal 3.1", "Pengecualian umum: kelalaian/kesengajaan tertanggung"),
    ("Pasal 3.2", "Cacat tersembunyi, keausan biasa, atau kerusakan perlahan-lahan"),
    ("Pasal 3.3", "Kerugian pada ban/velg akibat pemakaian biasa"),
    ("Pasal 3.4", "Kegagalan mekanis, kerusakan mesin akibat masuk air (bukan banjir)"),
    ("Pasal 3.5", "Pemakaian kendaraan melebihi kapasitas yang diizinkan"),
    ("Pasal 3.6", "Kerugian akibat barang bawaan/aksesoris tidak tercantum dalam polis"),
    ("Pasal 4.1", "Perang, invasi, aksi militer, pemberontakan"),
    ("Pasal 4.2", "Pemogokan, huru-hara, kerusuhan, terorisme (tanpa perluasan RSCC)"),
    ("Pasal 4.3", "Bencana alam gempa bumi, letusan gunung berapi, tsunami"),
    ("Pasal 4.4", "Banjir, genangan air, badai (tanpa perluasan banjir)"),
    ("Pasal 5.1", "Kendaraan digunakan di luar wilayah yang tercantum dalam polis"),
    ("Pasal 5.2", "Pengemudi tidak memiliki SIM yang sah sesuai golongan kendaraan"),
    ("Pasal 5.3", "Pengemudi dalam pengaruh alkohol, narkoba, atau obat-obatan terlarang"),
    ("Pasal 5.4", "Kendaraan digunakan untuk tindak kejahatan"),
    ("Pasal 5.5", "Kendaraan digunakan untuk balapan/uji kecepatan"),
    ("Pasal 6.1", "Laporan klaim tidak sesuai fakta (material misrepresentation/fraud)"),
    ("Pasal 6.2", "Klaim diajukan melewati batas waktu yang ditentukan polis"),
    ("Pasal 7.1", "Modifikasi tanpa persetujuan penanggung"),
    ("Pasal 8.1", "Kerugian akibat reaksi nuklir, radiasi, atau kontaminasi radioaktif"),
    ("Pasal 9.1", "Kerugian yang telah dijamin oleh asuransi lain (double insurance)"),
]


# ──────────────────────────────────────────────────────────────
# HELPER
# ──────────────────────────────────────────────────────────────

def _contains_drawing(p_elem) -> bool:
    return p_elem.find('.//' + qn('w:drawing')) is not None

def _is_in_fallback(elem) -> bool:
    p = elem.getparent()
    while p is not None:
        if p.tag == f"{{{NS_MC}}}Fallback":
            return True
        p = p.getparent()
    return False

def _para_text(p) -> str:
    """Teks bersih paragraf, skip fallback."""
    parts = []
    for node in p._p.iter(qn('w:t')):
        if not _is_in_fallback(node):
            parts.append(node.text or "")
    return "".join(parts)


# ──────────────────────────────────────────────────────────────
# FIX 2: NEWLINE → <w:br/> di Word
# ──────────────────────────────────────────────────────────────

def _inject_line_breaks(p_elem, text: str):
    """
    Ganti teks di paragraf yang mengandung '\\n' dengan multiple runs + <w:br/>.
    Ini untuk keterangan saksi dan kronologis yang diinput user dengan Enter.

    Word tidak mengenal '\\n' di w:t → harus pakai <w:br/> (soft return)
    yang berarti baris baru dalam paragraf yang sama (formatting tetap sama).

    Strategi:
    - Ambil formatting (rPr) dari run pertama yang punya konten
    - Hapus semua runs yang ada
    - Buat runs baru: satu per baris, diselingi <w:br/> antar baris
    """
    if '\n' not in text:
        return  # Tidak ada newline, tidak perlu diproses

    lines = text.split('\n')

    # Ambil rPr dari run pertama yang ada konten (untuk copy formatting)
    template_rPr = None
    for r_elem in p_elem.findall('.//' + qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is not None:
            template_rPr = copy.deepcopy(rPr)
            break

    # Hapus semua w:r yang ada di paragraf ini
    for r_elem in list(p_elem.findall('.//' + qn('w:r'))):
        parent = r_elem.getparent()
        if parent is not None:
            parent.remove(r_elem)

    # Buat runs baru: teks baris + <w:br/> sebagai pemisah
    for i, line in enumerate(lines):
        # Run teks
        r = OxmlElement('w:r')
        if template_rPr is not None:
            r.append(copy.deepcopy(template_rPr))
        wt = OxmlElement('w:t')
        wt.text = line
        if line and (line[0] == ' ' or line[-1] == ' '):
            wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(wt)
        p_elem.append(r)

        # <w:br/> setelah setiap baris kecuali yang terakhir
        if i < len(lines) - 1:
            r_br = OxmlElement('w:r')
            if template_rPr is not None:
                r_br.append(copy.deepcopy(template_rPr))
            br = OxmlElement('w:br')
            r_br.append(br)
            p_elem.append(r_br)


# ──────────────────────────────────────────────────────────────
# CORE REPLACE (v3 proven)
# ──────────────────────────────────────────────────────────────

# Placeholder yang nilainya mungkin mengandung newline
_NEWLINE_PLACEHOLDERS = {
    '[keterangan_saksi1]', '[keterangan_saksi2]', '[keterangan_saksi3]',
    '[keterangan_saksi4]', '[keterangan_saksi5]', '[keterangan_saksi6]',
    '[keterangan_saksi7]', '[keterangan_saksi8]', '[keterangan_saksi9]',
    '[keterangan_saksi10]', '[kronologis_kejadian]', '[keterangan_tkp]',
}


def _replace_in_paragraph(p_elem, replacements: dict):
    """Replace placeholder. Jika nilai mengandung newline → inject w:br."""
    wt_nodes = [n for n in p_elem.iter(qn('w:t'))]
    if not wt_nodes:
        return

    # Cek apakah paragraf ini mengandung placeholder yg nilainya punya newline
    full_before = "".join(n.text or "" for n in wt_nodes)

    # Lapis 1: replace per w:t
    for wt in wt_nodes:
        if not wt.text:
            continue
        orig = wt.text
        new = orig
        for ph, val in replacements.items():
            new = new.replace(ph, str(val) if val is not None else "")
        if new != orig:
            wt.text = new
            if new and (new[0] == ' ' or new[-1] == ' '):
                wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Re-scan
    full_after = "".join(n.text or "" for n in wt_nodes)
    has_split = any(ph in full_after for ph in replacements)

    # Lapis 2: hanya jika masih ada placeholder terpecah
    if has_split:
        new_combined = full_after
        for ph, val in replacements.items():
            new_combined = new_combined.replace(ph, str(val) if val is not None else "")
        if new_combined != full_after:
            target = next((wt for wt in wt_nodes if wt.text and wt.text.strip()), wt_nodes[0])
            written = False
            for wt in wt_nodes:
                if wt is target:
                    wt.text = new_combined
                    if new_combined and (new_combined[0] == ' ' or new_combined[-1] == ' '):
                        wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    written = True
                elif written and wt.text and wt.text.strip():
                    wt.text = ""
            full_after = new_combined

    # FIX 2: Cek apakah teks hasil replace mengandung newline
    # Ini terjadi ketika [keterangan_saksiN] atau [kronologis_kejadian] punya '\n'
    if '\n' in full_after:
        _inject_line_breaks(p_elem, full_after)


def _replace_in_textboxes(container_elem, replacements: dict):
    for choice in container_elem.iter(f"{{{NS_MC}}}Choice"):
        for txbx in choice.iter(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'
        ):
            for p_elem in txbx.findall(qn('w:p')):
                _replace_in_paragraph(p_elem, replacements)


def _replace_all(doc: Document, replacements: dict):
    def _do_paragraphs(paragraphs):
        for p in paragraphs:
            if _contains_drawing(p._p):
                continue
            _replace_in_paragraph(p._p, replacements)

    def _do_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _do_paragraphs(cell.paragraphs)
                    _do_tables(cell.tables)

    _do_paragraphs(doc.paragraphs)
    _do_tables(doc.tables)
    _replace_in_textboxes(doc.element.body, replacements)

    for section in doc.sections:
        for part in [
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ]:
            if part:
                _do_paragraphs(part.paragraphs)
                _do_tables(part.tables)
                if hasattr(part, '_element'):
                    _replace_in_textboxes(part._element, replacements)


# ──────────────────────────────────────────────────────────────
# FIX 1 & 3: HAPUS BLOK TIDAK TERPAKAI (PRESISI, JAGA SPACING)
# ──────────────────────────────────────────────────────────────


def _trim_saksi_to_data_spacing(doc: Document):
    """
    Potong paragraf kosong berlebih antara saksi terakhir dan DATA DIPEROLEH.
    Sisakan hanya 2 empty (= 2 enter). Akses XML langsung agar tidak stale.
    """
    body = doc.element.body
    # Ambil semua w:p langsung dari body (bukan rekursif — hanya body level)
    all_p = [child for child in body if child.tag == qn('w:p')]

    # Temukan index DATA DIPEROLEH
    data_idx = None
    for i, p in enumerate(all_p):
        t = "".join(n.text or "" for n in p.iter(qn('w:t')))
        if 'DATA' in t and 'DIPEROLEH' in t:
            data_idx = i
            break
    if data_idx is None:
        return

    # Kumpulkan empty berurutan tepat sebelum DATA (scan mundur)
    empty_elems = []
    for i in range(data_idx - 1, -1, -1):
        t = "".join(n.text or "" for n in all_p[i].iter(qn('w:t'))).strip()
        if t == '':
            empty_elems.insert(0, all_p[i])
        else:
            break

    # Hapus kelebihan — sisakan 2 terakhir
    MAX_EMPTY = 2
    to_delete = empty_elems[:-MAX_EMPTY] if len(empty_elems) > MAX_EMPTY else []
    for p_elem in to_delete:
        if p_elem in list(body):
            body.remove(p_elem)

def _remove_unused_blocks(doc: Document, jumlah_saksi: int, jumlah_analisa: int):
    """
    Hapus paragraf dari blok saksi/analisa yang tidak terpakai.
    Prinsip: HANYA hapus paragraf yang merupakan konten blok kosong.
    Paragraf spacer/empty ANTARA blok yang terpakai → JANGAN disentuh.

    ANALISA (FIX 1):
    Pola template: [analisaN] → '' → [analisaN+1]
    Setelah replace analisa tidak terpakai: '' (kosong) → '' (spacer) → ...
    Fix: hapus pasangan (baris_analisa_kosong + spacer_setelahnya) sekaligus.
    Deteksi: paragraf yang setelah replace = '' DAN paragraf tepat sebelumnya
    adalah [analisaN] yang sudah menjadi kosong.

    SAKSI (FIX 3):
    Tiap blok detail saksi = [nama+role, alamat, '', usia/tgl/pekerjaan, '', keterangan, '']
    Untuk saksi tidak terpakai, setelah replace semua jadi:
      '( )' atau '(  )' → nama+role kosong
      'Alamat :' atau 'Alamat : ' → alamat kosong
      '' (empty) → spacer
      'Saksi (th) ..., , ...' → usia/tanggal/pekerjaan kosong
      '' (empty) → spacer
      '' (empty) → keterangan kosong
      '' (empty) → spacer setelah blok
    Fix: hapus semua 7 paragraf ini sekaligus sebagai satu unit.
    Baris singkat saksi ("( )") di section list juga dihapus.
    Baris "Hasil interview dari  " juga dihapus.
    Spacer ANTARA blok yang terpakai → TIDAK disentuh.
    """
    paras = doc.paragraphs
    to_remove_elems = set()  # set of id(p._p)

    # ── ANALISA: hapus pasangan (analisa_kosong + spacer)
    # Identifikasi paragraf analisa berdasarkan index (analisa ke-N+1 sampai 10)
    # Setelah replace mereka sudah jadi '' atau '.' — ambil berdasarkan posisi relatif
    # terhadap marker "ANALISA FAKTA"

    analisa_idx = []  # index paragraf [analisaN] di template, berurutan
    in_analisa = False
    for i, p in enumerate(paras):
        t = _para_text(p).strip()
        if 'ANALISA FAKTA' in t:
            in_analisa = True
            continue
        if in_analisa and ('Maka dari' in t or 'KRONOLOGI' in t or 'Kronologi' in t):
            in_analisa = False
        if in_analisa:
            # Paragraf analisa: yang di template ber-numPr (numbered list)
            # Setelah replace, yang tidak terpakai jadi '' atau '.'
            pPr = p._p.find(qn('w:pPr'))
            is_numbered = pPr is not None and pPr.find(qn('w:numPr')) is not None
            if is_numbered:
                analisa_idx.append(i)

    # Hapus analisa ke-(jumlah_analisa+1) sampai terakhir beserta spacer-nya
    for i in analisa_idx[jumlah_analisa:]:
        p_analisa = paras[i]
        t = _para_text(p_analisa).strip()
        # Hanya hapus jika memang kosong (sudah replace dengan "")
        if t in ('', '.'):
            to_remove_elems.add(id(p_analisa._p))
            # Hapus spacer tepat sesudahnya jika ada dan kosong
            if i + 1 < len(paras):
                p_next = paras[i + 1]
                t_next = _para_text(p_next).strip()
                if t_next == '':
                    to_remove_elems.add(id(p_next._p))

    # ── SAKSI LIST SINGKAT: hapus baris "( )" yang kosong
    in_list_zone = False
    for i, p in enumerate(paras):
        t = _para_text(p)
        t_s = t.strip()
        if 'Saksi - Saksi yang diinterview' in t or 'Saksi-Saksi yang diinterview' in t:
            in_list_zone = True
            continue
        if in_list_zone and ('Keterangan Saksi' in t):
            in_list_zone = False
        if in_list_zone:
            # Baris list saksi tidak terpakai: "( )" atau "(  )" atau "(   )"
            if re.match(r'^\([\s]*\)[\s]*$', t_s) and t_s != '':
                to_remove_elems.add(id(p._p))
            elif t_s == '( )' or t_s == '(  )' or t_s == '()':
                to_remove_elems.add(id(p._p))

    # ── SAKSI DETAIL: hapus blok berdasarkan penanda "Saksi (th)" dan "Alamat : " kosong
    # Cara paling reliable: scan seluruh zona, tandai SETIAP paragraf yang
    # merupakan sisa blok saksi tidak terpakai berdasarkan isinya setelah replace:
    #   - "Saksi (th) ..." tanpa angka usia → usia kosong = blok tidak terpakai
    #   - "Alamat : " tanpa nilai → alamat kosong = blok tidak terpakai
    #   - baris yang hanya berisi "( )" atau "( )  " trailing spaces → nama+role kosong
    #   - keterangan kosong ("") tepat setelah baris usia kosong → keterangan blok ini
    #   - spacer "" setelah keterangan kosong → spacer blok ini

    in_detail_zone = False
    i = 0
    paras_snapshot = list(paras)

    while i < len(paras_snapshot):
        t = _para_text(paras_snapshot[i]).strip()
        t_raw = _para_text(paras_snapshot[i])

        if 'Keterangan Saksi' in t_raw:
            in_detail_zone = True
            i += 1
            continue
        if in_detail_zone and 'DATA' in t and 'DIPEROLEH' in t:
            in_detail_zone = False

        if in_detail_zone:
            # Deteksi "Saksi (th)" tanpa angka usia → blok tidak terpakai
            is_empty_usia = bool(re.match(r'^Saksi \(th\)', t))
            # Deteksi "Alamat : " tanpa nilai
            is_empty_alamat = bool(re.match(r'^Alamat\s*:\s*$', t))
            # Deteksi nama+role: pola "...( ... )  " — cek apakah hanya "( )" dengan trailing
            is_empty_name = bool(re.match(r'^\(\s*\)\s*$', t))

            if is_empty_usia or is_empty_alamat or is_empty_name:
                to_remove_elems.add(id(paras_snapshot[i]._p))

        i += 1

    # Tambahan: hapus baris keterangan kosong yang LANGSUNG setelah baris usia kosong
    # (karena baris usia kosong ditandai tapi baris keterangan kosong belum)
    # Scan sekali lagi untuk empty yang diapit oleh baris yang sudah di-mark
    i = 1
    paras_snapshot2 = list(paras)
    while i < len(paras_snapshot2) - 1:
        prev_id = id(paras_snapshot2[i-1]._p)
        curr_t  = _para_text(paras_snapshot2[i]).strip()
        next_id = id(paras_snapshot2[i+1]._p) if i+1 < len(paras_snapshot2) else None

        if prev_id in to_remove_elems and curr_t == '':
            to_remove_elems.add(id(paras_snapshot2[i]._p))
        i += 1

    # ── "Hasil interview dari  " kosong
    in_data_zone = False
    for p in paras:
        t = _para_text(p)
        t_s = t.strip()
        if 'DATA' in t and 'DIPEROLEH' in t:
            in_data_zone = True
            continue
        if in_data_zone and 'ANALISA' in t:
            in_data_zone = False
        if in_data_zone:
            if re.match(r'^Hasil interview dari\s*$', t_s):
                to_remove_elems.add(id(p._p))

    # ── DEDUP: hapus baris nama+role yang muncul duplikat di zona saksi
    # Terjadi karena typo di template (misal [saksi3] dipakai dua kali)
    in_det = False
    seen_names: dict = {}
    for i, p in enumerate(paras):
        t = _para_text(p)
        t_s = t.strip()
        if 'Keterangan Saksi' in t:
            in_det = True
        if in_det and 'IV. DATA' in t:
            in_det = False
        if in_det:
            # Baris nama+role = ada "(" dan ")" tapi bukan baris usia/alamat
            if ('(' in t_s and ')' in t_s
                    and not t_s.startswith('Saksi')
                    and not t_s.startswith('Alamat')):
                if t_s in seen_names:
                    # Duplikat → hapus beserta empty langsung sesudahnya
                    to_remove_elems.add(id(p._p))
                    if i + 1 < len(paras):
                        if _para_text(paras[i + 1]).strip() == '':
                            to_remove_elems.add(id(paras[i + 1]._p))
                else:
                    seen_names[t_s] = i

    # ── Eksekusi hapus (batch)
    for p in paras:
        if id(p._p) in to_remove_elems:
            parent = p._p.getparent()
            if parent is not None and p._p in list(parent):
                parent.remove(p._p)

    # ── TERAKHIR: trim empty berlebih sebelum DATA DIPEROLEH → sisakan 2 saja
    # (dilakukan setelah hapus blok agar hasilnya bersih)
    _trim_saksi_to_data_spacing(doc)


# ──────────────────────────────────────────────────────────────
# FIX 4: TTD SURVEYOR — insert gambar tanda tangan
# ──────────────────────────────────────────────────────────────

def _insert_ttd_surveyor(doc: Document, surveyor_name: str, ttd_dir: str = 'ttd'):
    """
    Insert TTD surveyor sebagai inline image di paragraf kosong
    yang tepat berada 1 baris DI ATAS baris nama direktur+surveyor.

    Template area TTD:
      "     Mengetahui,"
      "" (kosong x5)
      "Dharmawan Gunawan, S.H.  NamaSurveyor"   ← nama
      "     Direktur Utama      Surveyor"

    Kita scan mundur dari baris nama → cari paragraf kosong pertama
    yang ditemukan → itu target insert TTD.
    Gambar disisipkan sebagai inline run + diawali tab untuk
    mendorong ke posisi kolom kanan (sejajar surveyor).
    """
    from docx.shared import Cm

    # Resolve file TTD
    ttd_path = None
    for ext in ['png', 'jpg', 'jpeg']:
        p1 = os.path.join(ttd_dir, f"{surveyor_name}.{ext}")
        if os.path.exists(p1):
            ttd_path = p1; break
    if ttd_path is None:
        alt = surveyor_name.replace(' ', '_').replace('.', '').lower()
        for ext in ['png', 'jpg', 'jpeg']:
            p2 = os.path.join(ttd_dir, f"{alt}.{ext}")
            if os.path.exists(p2):
                ttd_path = p2; break
    if ttd_path is None:
        return  # Tidak ada file → skip

    # Temukan baris nama direktur+surveyor
    nama_idx = None
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        t = _para_text(p)
        if surveyor_name in t and 'S.H.' in t:
            nama_idx = i
            break
    if nama_idx is None:
        return

    # Scan mundur dari nama_idx, cari paragraf kosong pertama
    # (langsung di atas nama = posisi TTD yang benar)
    target_para = None
    for i in range(nama_idx - 1, max(0, nama_idx - 8), -1):
        t = _para_text(paras[i]).strip()
        if t == '':
            target_para = paras[i]
            break
        else:
            break  # Ada teks → tidak ada empty di atas nama

    if target_para is None:
        return

    try:
        # Tambah run tab dulu (mendorong ke kolom kanan sejajar surveyor)
        r_tab = OxmlElement('w:r')
        w_tab_elem = OxmlElement('w:tab')
        r_tab.append(w_tab_elem)

        # Tambah run gambar via python-docx (handle rId otomatis)
        run_pic = target_para.add_run()
        run_pic.add_picture(ttd_path, width=Cm(2.5))

        # Ambil w:r dari run_pic dan pindahkan ke posisi setelah tab
        pic_r = target_para._p.findall(qn('w:r'))[-1]
        target_para._p.remove(pic_r)

        # Insert tab + gambar di awal paragraf (setelah w:pPr)
        pPr = target_para._p.find(qn('w:pPr'))
        base_idx = (list(target_para._p).index(pPr) + 1) if pPr is not None else 0
        target_para._p.insert(base_idx, r_tab)
        target_para._p.insert(base_idx + 1, pic_r)

    except Exception:
        pass  # Gagal insert gambar → skip tanpa crash


def _build_replacements(data: dict) -> dict:
    R = {}

    asuransi    = data.get('asuransi', '')
    tertanggung = data.get('tertanggung', '')
    nopolis     = data.get('nopolis', '')
    merk        = data.get('merk', '')
    nopol       = data.get('nopol', '')
    case        = data.get('case', '')

    R['[ASURANSI]']    = asuransi.upper()
    R['[TERTANGGUNG]'] = tertanggung.upper()
    R['[NOPOLIS]']     = nopolis.upper()
    R['[MERK]']        = merk.upper()
    R['[NOPOL]']       = nopol.upper()
    R['[CASE]']        = case.upper()

    R['[asuransi]']    = asuransi.title()
    R['[tertanggung]'] = tertanggung
    R['[nopolis]']     = nopolis.upper()
    R['[merk]']        = merk.title()
    R['[nopol]']       = nopol.upper()
    R['[case]']        = case.lower()

    R['[surat_kuasa]']         = data.get('surat_kuasa', '')
    R['[tanggal_surat_kuasa]'] = data.get('tanggal_surat_kuasa', '')
    R['[surat_tugas]']         = data.get('surat_tugas', '')
    R['[tanggal_surat_tugas]'] = data.get('tanggal_surat_tugas', '')
    R['[surveyor]']            = data.get('surveyor', '')
    R['[alamat_tertanggung]']  = data.get('alamat_tertanggung', '')
    R['[periode]']             = data.get('periode', '')
    R['[coverage]']            = data.get('coverage', '')
    R['[usage]']               = data.get('usage', '')
    R['[tahun_kendaraan]']     = data.get('tahun_kendaraan', '')
    R['[no_ka]']               = data.get('no_ka', '')
    R['[no_sin]']              = data.get('no_sin', '')
    R['[penyebab]']            = data.get('penyebab', '')
    R['[dol]']                 = data.get('dol', '')
    R['[dol_time]']            = data.get('dol_time', '')
    R['[alamat_tkp]']          = data.get('alamat_tkp', '')
    R['[keterangan_tkp]']      = data.get('keterangan_tkp', '')
    R['[dol_lhs]']             = data.get('dol_lhs', '')
    R['[kronologis_kejadian]'] = data.get('kronologis_kejadian', '')

    claimable = data.get('claimable', True)
    R['[ditemukan_atau_tidak]'] = 'tidak ditemukan' if claimable else 'ditemukan'
    R['[dapat_atau_tidak]']     = 'dapat' if claimable else 'tidak dapat'
    if claimable:
        R['[claimable_atau_tidak]'] = 'CLAIMABLE'
    else:
        pasal = data.get('pasal_unclaimable', '')
        R['[claimable_atau_tidak]'] = f'UNCLAIMABLE - {pasal}' if pasal else 'UNCLAIMABLE'

    saksi_list = data.get('saksi_list', [])
    for i in range(1, 11):
        if i <= len(saksi_list):
            s = saksi_list[i - 1]
            # Normalisasi newline: \r\n → \n
            ket = s.get('keterangan', '').replace('\r\n', '\n').replace('\r', '\n')
            R[f'[saksi{i}]']            = s.get('nama', '')
            R[f'[role{i}]']             = s.get('role', '')
            R[f'[alamat_saksi{i}]']     = s.get('alamat', '')
            R[f'[usia{i}]']             = s.get('usia', '')
            R[f'[dol_interview{i}]']    = s.get('tanggal_interview', '')
            R[f'[pekerjaan{i}]']        = s.get('pekerjaan', '')
            R[f'[keterangan_saksi{i}]'] = ket
        else:
            R[f'[saksi{i}]']            = ''
            R[f'[role{i}]']             = ''
            R[f'[alamat_saksi{i}]']     = ''
            R[f'[usia{i}]']             = ''
            R[f'[dol_interview{i}]']    = ''
            R[f'[pekerjaan{i}]']        = ''
            R[f'[keterangan_saksi{i}]'] = ''

    analisa_list = data.get('analisa_list', [])
    for i in range(1, 11):
        R[f'[analisa{i}]'] = analisa_list[i - 1] if i <= len(analisa_list) else ''

    # Normalisasi newline kronologis
    kron = data.get('kronologis_kejadian', '').replace('\r\n', '\n').replace('\r', '\n')
    R['[kronologis_kejadian]'] = kron

    return R


# ──────────────────────────────────────────────────────────────
# PUBLIC API
# ──────────────────────────────────────────────────────────────

def generate_lhs(template_path: str, output_path: str, data: dict,
                 ttd_dir: str = 'ttd') -> str:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template tidak ditemukan: {template_path}")

    doc = Document(template_path)

    # Step 1: Replace semua placeholder
    replacements = _build_replacements(data)
    _replace_all(doc, replacements)

    # Step 2: Hapus blok tidak terpakai (saksi & analisa) — presisi, jaga spacing
    jumlah_saksi   = len([s for s in data.get('saksi_list', []) if s.get('nama', '').strip()])
    jumlah_analisa = len([a for a in data.get('analisa_list', []) if str(a).strip()])
    _remove_unused_blocks(doc, jumlah_saksi, jumlah_analisa)

    # Step 3: Insert TTD surveyor jika file tersedia
    surveyor = data.get('surveyor', '')
    if surveyor and ttd_dir:
        _insert_ttd_surveyor(doc, surveyor, ttd_dir)

    doc.save(output_path)
    return output_path